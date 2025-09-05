from docx import Document
from docx.shared import Inches
import os

# List your image filenames here
image_files = [
    "6ce42003-e0c2-423a-ac13-343050d8a729.png",
    "boxplot(new).png",
    "Plot 1.png",
    "PLot 2.png",
    "Plot 3(new).png",
    "Plot 3.png",
    "Plot 4.png",
    "Plot 5.png",
    "Plot 6.png",
    "Plot 7(new).png"
]

# Create a Word document
doc = Document()
doc.add_heading("Visual Data Analysis", level=1)

# Add each image with a caption
for idx, image in enumerate(image_files, start=1):
    doc.add_paragraph(f"Figure {idx}:", style='Heading 2')
    if os.path.exists(image):
        doc.add_picture(image, width=Inches(6))
        doc.add_paragraph("")  # Spacing
    else:
        doc.add_paragraph(f"[Missing file: {image}]")

# Save the document
output_path = "Data_Visualizations.docx"
doc.save(output_path)
print(f"Document saved as: {output_path}")
